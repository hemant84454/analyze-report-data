#!/usr/local/bin/python3.6

import logging
import pandas as pd

from docx.api import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.text.run import Font, Run
from docx.shared import RGBColor


def sortResults(results_items):
	column_headers = ['table_type', 'table_title', 'color', 'message']
	results_df = pd.DataFrame(columns=column_headers)

	def sort_color(color):
		try:
			if str(color).lower() == 'green':
				return 1
			elif str(color).lower() == 'yellow':
				return 2
			elif str(color).lower() == 'red':
				return 3
			else:
				return 4
		except:
			return 5

	for item in results_items:
		results_df = results_df.append({'table_type': item['table_type'], 'table_title': item['table_title'], 'color': item['color'], 'message': item['message']}, ignore_index=True)

	results_df['sort_color'] = results_df['color'].apply(lambda x: sort_color(x))
	results_df.sort_values(by=['table_type', 'table_title', 'sort_color', 'message'], inplace=True)
	# print(results_df.to_string())

	return results_df

def createDocxFile(output_file, file_name, analysis_type, results):
	document = Document('output_template.docx')

	# document.add_picture('RedThread.png', width=Inches(2))
	# last_paragraph = document.paragraphs[-1]
	# last_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
	try:
		if analysis_type is not None:
			analysis_type = analysis_type.lower()
	except:
		analysis_type = ''

	if 'lm' in analysis_type:
		first_part = 'Large Molecule'
	elif 'sm' in analysis_type:
		first_part = 'Small Molecule'
	else:
		first_part = ''

	if 'mv' in analysis_type:
		second_part = ' Method Validation Findings'
	elif 'sm' in analysis_type:
		second_part = ' Sample Analysis Findings'
	else:
		second_part = ''


	document.add_heading(first_part+second_part, 0)
	last_paragraph = document.paragraphs[-1]
	last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

	document.add_heading(str(file_name), 1)
	last_paragraph = document.paragraphs[-1]
	last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

	prev_table_type = ''
	prev_table_title = ''
	prev_color = ''

	for row in results:
		table_type = str(row['heading']).replace('\n', ' ').replace('\r', ' ').strip()
		document.add_heading(table_type, 2)

		table_type_items = row['resultItem']
		for table_type_item in table_type_items:
			table_title = str(table_type_item['tabletype']).replace('\n', ' ').replace('\r', ' ').strip()
			p = document.add_paragraph()
			run = p.add_run(table_title)
			run.font.bold = True

			table_title_items = table_type_item['item']
			for table_title_item in table_title_items:
				color = str(table_title_item['color']).replace('\n', ' ').replace('\r', ' ').strip()
				p2 = document.add_paragraph()
				run2 = p2.add_run()
				run2.add_text(color)
				run2.font.underline = True

				color_items = table_title_item['message']
				for color_item in color_items:
					message = str(color_item).replace('\n', ' ').replace('\r', ' ').strip()
					p3 = document.add_paragraph(message, style="List Bullet 2")
					p3_format = p3.paragraph_format

		p = document.add_paragraph()

	document.save(output_file)
